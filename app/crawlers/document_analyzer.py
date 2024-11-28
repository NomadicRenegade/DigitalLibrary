import fitz  # PyMuPDF
import docx
import openpyxl
from typing import Dict, Any
from pathlib import Path

class DocumentAnalyzer:
    """Analyzes document files (PDF, DOCX, XLSX) for metadata and content"""
    
    @staticmethod
    def analyze_pdf(filepath: str) -> Dict[str, Any]:
        """Extract metadata and content from PDF files"""
        doc = fitz.open(filepath)
        
        metadata = {
            'page_count': len(doc),
            'title': doc.metadata.get('title', ''),
            'author': doc.metadata.get('author', ''),
            'subject': doc.metadata.get('subject', ''),
            'keywords': doc.metadata.get('keywords', ''),
            'text_content': '',
            'has_images': False,
            'image_count': 0
        }
        
        # Extract text and check for images
        image_count = 0
        text_content = []
        
        for page in doc:
            text_content.append(page.get_text())
            image_list = page.get_images()
            image_count += len(image_list)
            
        metadata['text_content'] = '\n'.join(text_content)
        metadata['has_images'] = image_count > 0
        metadata['image_count'] = image_count
        
        return metadata
        
    @staticmethod
    def analyze_docx(filepath: str) -> Dict[str, Any]:
        """Extract metadata and content from DOCX files"""
        doc = docx.Document(filepath)
        
        metadata = {
            'paragraph_count': len(doc.paragraphs),
            'text_content': '',
            'has_tables': bool(doc.tables),
            'table_count': len(doc.tables)
        }
        
        # Extract text content
        text_content = []
        for paragraph in doc.paragraphs:
            text_content.append(paragraph.text)
            
        metadata['text_content'] = '\n'.join(text_content)
        
        return metadata
        
    @staticmethod
    def analyze_xlsx(filepath: str) -> Dict[str, Any]:
        """Extract metadata and content from XLSX files"""
        workbook = openpyxl.load_workbook(filepath, read_only=True)
        
        metadata = {
            'sheet_count': len(workbook.sheetnames),
            'sheet_names': workbook.sheetnames,
            'has_charts': False,
            'has_formulas': False,
            'cell_count': 0
        }
        
        # Analyze sheets
        for sheet_name in workbook.sheetnames:
            sheet = workbook[sheet_name]
            cell_count = 0
            
            for row in sheet.iter_rows():
                for cell in row:
                    if cell.value is not None:
                        cell_count += 1
                        if isinstance(cell.value, str) and cell.value.startswith('='):
                            metadata['has_formulas'] = True
                            
            metadata['cell_count'] += cell_count
            
        return metadata
        
    def analyze(self, filepath: str) -> Dict[str, Any]:
        """Analyze document based on file type"""
        path = Path(filepath)
        
        try:
            if path.suffix.lower() == '.pdf':
                return self.analyze_pdf(filepath)
            elif path.suffix.lower() == '.docx':
                return self.analyze_docx(filepath)
            elif path.suffix.lower() in ['.xlsx', '.xls']:
                return self.analyze_xlsx(filepath)
            else:
                raise ValueError(f"Unsupported file type: {path.suffix}")
        except Exception as e:
            return {
                'error': str(e),
                'status': 'failed'
            }